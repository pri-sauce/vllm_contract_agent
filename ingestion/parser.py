# """
# ingestion/parser.py — Document ingestion layer
# Handles PDF, DOCX, and plain text contracts.
# Extracts clean text while preserving structure.
# """

# from pathlib import Path
# from dataclasses import dataclass, field
# from typing import Optional
# import re

# from loguru import logger


# @dataclass
# class ParsedDocument:
#     """
#     Represents a fully parsed contract document.
#     This is the standard format passed through the entire pipeline.
#     """
#     filename: str
#     file_type: str                      # pdf | docx | txt
#     raw_text: str                       # Full extracted text
#     pages: list[str] = field(default_factory=list)   # Text per page
#     metadata: dict = field(default_factory=dict)     # Title, author, etc.
#     word_count: int = 0
#     char_count: int = 0

#     def __post_init__(self):
#         self.word_count = len(self.raw_text.split())
#         self.char_count = len(self.raw_text)


# class DocumentParser:
#     """
#     Parses contract documents into clean text.
#     Supports: PDF (native + scanned), DOCX, TXT
#     """

#     def parse(self, file_path: str | Path) -> ParsedDocument:
#         """
#         Main entry point. Auto-detects file type and routes to correct parser.
#         """
#         path = Path(file_path)

#         if not path.exists():
#             raise FileNotFoundError(f"File not found: {path}")

#         suffix = path.suffix.lower()

#         logger.info(f"Parsing {path.name} ({suffix})")

#         if suffix == ".pdf":
#             return self._parse_pdf(path)
#         elif suffix in (".docx", ".doc"):
#             return self._parse_docx(path)
#         elif suffix == ".txt":
#             return self._parse_txt(path)
#         else:
#             raise ValueError(f"Unsupported file type: {suffix}. Supported: PDF, DOCX, TXT")

#     # ------------------------------------------------------------------
#     # PDF Parser
#     # ------------------------------------------------------------------

#     def _parse_pdf(self, path: Path) -> ParsedDocument:
#         try:
#             import fitz  # PyMuPDF
#         except ImportError:
#             raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")

#         doc = fitz.open(str(path))
#         pages = []
#         metadata = {}

#         # Extract metadata
#         meta = doc.metadata
#         if meta:
#             metadata = {
#                 "title": meta.get("title", ""),
#                 "author": meta.get("author", ""),
#                 "subject": meta.get("subject", ""),
#                 "creator": meta.get("creator", ""),
#                 "page_count": len(doc),
#             }

#         for page_num, page in enumerate(doc):
#             text = page.get_text("text")

#             # If page has no text, it's likely scanned — try OCR
#             if not text.strip() and len(text.strip()) < 50:
#                 logger.warning(f"Page {page_num + 1} appears scanned. Attempting OCR...")
#                 text = self._ocr_page(page)

#             pages.append(self._clean_text(text))

#         doc.close()

#         full_text = "\n\n".join(pages)

#         return ParsedDocument(
#             filename=path.name,
#             file_type="pdf",
#             raw_text=full_text,
#             pages=pages,
#             metadata=metadata,
#         )

#     def _ocr_page(self, page) -> str:
#         """OCR fallback for scanned PDF pages using Tesseract."""
#         try:
#             import pytesseract
#             from PIL import Image
#             import io

#             # Render page to image at high resolution
#             mat = page.get_pixmap(dpi=300)
#             img_data = mat.tobytes("png")
#             img = Image.open(io.BytesIO(img_data))

#             text = pytesseract.image_to_string(img, lang="eng")
#             logger.info("OCR completed for page.")
#             return text

#         except ImportError:
#             logger.warning("pytesseract not installed. OCR skipped for scanned page.")
#             return ""
#         except Exception as e:
#             logger.error(f"OCR failed: {e}")
#             return ""

#     # ------------------------------------------------------------------
#     # DOCX Parser
#     # ------------------------------------------------------------------

#     def _parse_docx(self, path: Path) -> ParsedDocument:
#         try:
#             from docx import Document
#         except ImportError:
#             raise ImportError("python-docx not installed. Run: pip install python-docx")

#         doc = Document(str(path))
#         paragraphs = []
#         metadata = {}

#         # Extract core properties if available
#         try:
#             props = doc.core_properties
#             metadata = {
#                 "title": props.title or "",
#                 "author": props.author or "",
#                 "created": str(props.created) if props.created else "",
#             }
#         except Exception:
#             pass

#         for para in doc.paragraphs:
#             if para.text.strip():
#                 paragraphs.append(para.text.strip())

#         # Also extract text from tables (contracts often have tables)
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
#                 if row_text:
#                     paragraphs.append(f"[TABLE ROW] {row_text}")

#         full_text = "\n".join(paragraphs)

#         return ParsedDocument(
#             filename=path.name,
#             file_type="docx",
#             raw_text=self._clean_text(full_text),
#             pages=[full_text],  # DOCX doesn't have native pages
#             metadata=metadata,
#         )

#     # ------------------------------------------------------------------
#     # Plain Text Parser
#     # ------------------------------------------------------------------

#     def _parse_txt(self, path: Path) -> ParsedDocument:
#         text = path.read_text(encoding="utf-8", errors="replace")

#         return ParsedDocument(
#             filename=path.name,
#             file_type="txt",
#             raw_text=self._clean_text(text),
#             pages=[text],
#             metadata={},
#         )

#     # ------------------------------------------------------------------
#     # Text Cleaning
#     # ------------------------------------------------------------------

#     def _clean_text(self, text: str) -> str:
#         """
#         Clean extracted text while preserving contract structure.
#         Contracts rely heavily on numbering and indentation — be conservative.
#         """
#         # Normalize unicode characters
#         text = text.replace("\u2019", "'").replace("\u2018", "'")
#         text = text.replace("\u201c", '"').replace("\u201d", '"')
#         text = text.replace("\u2013", "-").replace("\u2014", "-")
#         text = text.replace("\u00a0", " ")  # Non-breaking space

#         # Remove excessive blank lines (keep max 2)
#         text = re.sub(r"\n{3,}", "\n\n", text)

#         # Remove page header/footer artifacts (common in PDFs)
#         text = re.sub(r"^\s*Page \d+ of \d+\s*$", "", text, flags=re.MULTILINE)
#         text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)  # lone page numbers

#         # Normalize whitespace within lines
#         lines = [re.sub(r"[ \t]+", " ", line) for line in text.split("\n")]
#         text = "\n".join(lines)

#         return text.strip()


# # Singleton
# parser = DocumentParser()



"""
ingestion/parser.py — Document ingestion layer
Handles PDF, DOCX, and plain text contracts.
Extracts clean text while preserving structure.
"""

from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional
import re

from loguru import logger


@dataclass
class ParsedDocument:
    """
    Represents a fully parsed contract document.
    This is the standard format passed through the entire pipeline.
    """
    filename: str
    file_type: str                      # pdf | docx | txt
    raw_text: str                       # Full extracted text
    pages: list[str] = field(default_factory=list)   # Text per page
    metadata: dict = field(default_factory=dict)     # Title, author, etc.
    word_count: int = 0
    char_count: int = 0

    def __post_init__(self):
        self.word_count = len(self.raw_text.split())
        self.char_count = len(self.raw_text)


class DocumentParser:
    """
    Parses contract documents into clean text.
    Supports: PDF (native + scanned), DOCX, TXT
    """

    def parse(self, file_path: str | Path) -> ParsedDocument:
        """
        Main entry point. Auto-detects file type and routes to correct parser.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        suffix = path.suffix.lower()

        logger.info(f"Parsing {path.name} ({suffix})")

        if suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix in (".docx", ".doc"):
            return self._parse_docx(path)
        elif suffix == ".txt":
            return self._parse_txt(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}. Supported: PDF, DOCX, TXT")

    # ------------------------------------------------------------------
    # PDF Parser
    # ------------------------------------------------------------------

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")

        doc = fitz.open(str(path))
        pages = []
        metadata = {}

        # Extract metadata
        meta = doc.metadata
        if meta:
            metadata = {
                "title": meta.get("title", ""),
                "author": meta.get("author", ""),
                "subject": meta.get("subject", ""),
                "creator": meta.get("creator", ""),
                "page_count": len(doc),
            }

        for page_num, page in enumerate(doc):
            text = page.get_text("text")

            # If page has no text, it's likely scanned — try OCR
            if not text.strip() and len(text.strip()) < 50:
                logger.warning(f"Page {page_num + 1} appears scanned. Attempting OCR...")
                text = self._ocr_page(page)

            pages.append(self._clean_text(text))

        doc.close()

        full_text = "\n\n".join(pages)

        return ParsedDocument(
            filename=path.name,
            file_type="pdf",
            raw_text=full_text,
            pages=pages,
            metadata=metadata,
        )

    def _ocr_page(self, page) -> str:
        """OCR fallback for scanned PDF pages using Tesseract."""
        try:
            import pytesseract
            from PIL import Image
            import io

            # Render page to image at high resolution
            mat = page.get_pixmap(dpi=300)
            img_data = mat.tobytes("png")
            img = Image.open(io.BytesIO(img_data))

            text = pytesseract.image_to_string(img, lang="eng")
            logger.info("OCR completed for page.")
            return text

        except ImportError:
            logger.warning("pytesseract not installed. OCR skipped for scanned page.")
            return ""
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""

    # ------------------------------------------------------------------
    # DOCX Parser
    # ------------------------------------------------------------------

    def _parse_docx(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document(str(path))
        paragraphs = []
        metadata = {}

        # Extract core properties if available
        try:
            props = doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "created": str(props.created) if props.created else "",
            }
        except Exception:
            pass

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract text from tables (contracts often have tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(f"[TABLE ROW] {row_text}")

        full_text = "\n".join(paragraphs)

        return ParsedDocument(
            filename=path.name,
            file_type="docx",
            raw_text=self._clean_text(full_text),
            pages=[full_text],  # DOCX doesn't have native pages
            metadata=metadata,
        )

    # ------------------------------------------------------------------
    # Plain Text Parser
    # ------------------------------------------------------------------

    def _parse_txt(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="replace")

        return ParsedDocument(
            filename=path.name,
            file_type="txt",
            raw_text=self._clean_text(text),
            pages=[text],
            metadata={},
        )

    # ------------------------------------------------------------------
    # Text Cleaning
    # ------------------------------------------------------------------


    # Common clause-start patterns — used to inject blank lines before each clause
    _CLAUSE_BREAK_RES = None  # initialised lazily on first use

    @classmethod
    def _get_break_patterns(cls):
        if cls._CLAUSE_BREAK_RES is None:
            cls._CLAUSE_BREAK_RES = [
                re.compile(r"(?<!" + "\n" + r")(?=\d{1,2}(?:\.\d{1,2}){0,3}[.)]\s+[A-Z])"),
                re.compile(r"(?<!" + "\n" + r")(?=(ARTICLE|SECTION)\s+(?:[IVX]+|\d+))", re.IGNORECASE),
                re.compile(r"(?<!" + "\n" + r")(?=(?:[A-Z][.)]\s|[(][a-z][)]\s)[A-Z])"),
                re.compile(r"(?<!" + "\n" + r")(?=(WHEREAS|RECITALS?|BACKGROUND|PREAMBLE))", re.IGNORECASE),
            ]
        return cls._CLAUSE_BREAK_RES

    def _clean_text(self, text: str) -> str:
        # Normalise unicode
        text = text.replace("\u2019", "'").replace("\u2018", "'")
        text = text.replace("\u201c", chr(34)).replace("\u201d", chr(34))
        text = text.replace("\u2013", "-").replace("\u2014", "-")
        text = text.replace("\u00a0", " ").replace("\u2022", "-")

        # Strip repeating page footer boilerplate (company name + address lines)
        import re as _re
        text = _re.sub(
            r"\n[A-Z][A-Za-z ]+(?:Private Limited|Pvt\.? Ltd\.?)\n"
            r"Address\s*:.*?(?=\n[A-Z\d(]|\Z)",
            "\n", text, flags=_re.DOTALL
        )
        text = _re.sub(r"^\s*Page \d+ of \d+\s*$", "", text, flags=_re.MULTILINE)
        text = _re.sub(r"^\s*\d+\s*$", "", text, flags=_re.MULTILINE)

        # Normalise whitespace within lines
        text = "\n".join(_re.sub(r"[ \t]+", " ", ln) for ln in text.split("\n"))

        # KEY FIX: inject blank line before every clause-start pattern
        for pat in self._get_break_patterns():
            text = pat.sub("\n\n", text)

        # Collapse 3+ blank lines to 2
        text = _re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()


# Singleton
parser = DocumentParser()